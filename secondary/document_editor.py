"""
Document Editor mode — a standalone Word-like editor (equation ribbon clone,
text formatting, open/print/export) with a right-hand AI panel that mirrors
modes 1-5 and writes generated content straight into the document.

Self-contained Flask Blueprint: owns its own DB connection and routes, same
pattern as bac_generator/routes.py. Registered in app.py with no extra wiring.
"""
import base64
import os
import re
import sqlite3
import uuid
from datetime import datetime
from html.parser import HTMLParser

from flask import Blueprint, request, redirect, render_template, session, jsonify, send_file
from werkzeug.utils import secure_filename
from markupsafe import escape
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

from secondary import ai, adjacent

document_editor_bp = Blueprint("document_editor", __name__)

UPLOAD_FOLDER = "uploads"
MODEL_EXTENSIONS = {"pdf", "docx", "txt"}
SCAN_EXTENSIONS = {"png", "jpg", "jpeg", "pdf"}

MODES = [
    {"key": "mode1", "icon": "💬", "title": "Asistent AI", "desc": "Conversație educațională.", "auto_insert": False},
    {"key": "mode2", "icon": "🧮", "title": "Limbaj Natural → Matematic", "desc": "Scrii exercițiul în limbaj normal și primești automat forma matematică corectă.", "auto_insert": True},
    {"key": "mode3", "icon": "📄", "title": "Creează după Model", "desc": "Încarci un model de test existent în format Word sau PDF.", "auto_insert": True},
    {"key": "mode4", "icon": "✍️", "title": "Transcriere test scris de mână", "desc": "Încarci o fișă scrisă de mână și o transformăm automat în format digital.", "auto_insert": True},
    {"key": "mode5", "icon": "🎓", "title": "Generator BAC", "desc": "Creează variante personalizate de BAC.", "auto_insert": True},
]
MODE_BY_KEY = {m["key"]: m for m in MODES}


# =========================================================
# DB
# =========================================================

def get_db():
    conn = sqlite3.connect("profu.db", check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def _require_login():
    return session.get("user_id") is not None


# =========================================================
# HELPERS
# =========================================================

def _file_read(filepath):
    extension = os.path.splitext(filepath)[1].lower()

    if extension == ".pdf":
        reader = PdfReader(filepath)
        text = ""
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text if text.strip() else "[PDF scanat - conținut va fi extras la generare]"

    if extension == ".docx":
        doc = DocxDocument(filepath)
        return "\n".join(p.text for p in doc.paragraphs)

    if extension == ".txt":
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()

    return "Unsupported file type"


def _allowed(filename, allowed_set):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in allowed_set


def _save_upload(document_id, mode_key, file_storage):
    filename = secure_filename(file_storage.filename)
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    stored_name = f"editor_{document_id}_{mode_key}_{filename}"
    path = os.path.join(UPLOAD_FOLDER, stored_name)
    file_storage.save(path)
    return path, filename


def _doc_or_none(conn, document_id, user_id):
    return conn.execute(
        "SELECT * FROM documents WHERE id = ? AND user_id = ?",
        (document_id, user_id)
    ).fetchone()


def _mode_conversation(conn, document_id, mode_key, user_id):
    return conn.execute(
        "SELECT * FROM conversations WHERE document_id = ? AND mode = ? AND user_id = ?",
        (document_id, mode_key, user_id)
    ).fetchone()


def _mode_messages(conn, conversation_id):
    return conn.execute(
        "SELECT * FROM messages WHERE conversation_id = ? ORDER BY id ASC",
        (conversation_id,)
    ).fetchall()


def _last_assistant_content(conn, conversation_id):
    row = conn.execute(
        "SELECT content FROM messages WHERE conversation_id = ? AND role = 'assistant' ORDER BY id DESC LIMIT 1",
        (conversation_id,)
    ).fetchone()
    return row["content"] if row else ""


def _last_user_content(conn, conversation_id):
    row = conn.execute(
        "SELECT content FROM messages WHERE conversation_id = ? AND role = 'user' ORDER BY id DESC LIMIT 1",
        (conversation_id,)
    ).fetchone()
    return row["content"] if row else ""


def _history_text(conn, conversation_id):
    rows = conn.execute(
        "SELECT role, content FROM messages WHERE conversation_id = ? AND role != 'system' ORDER BY id ASC",
        (conversation_id,)
    ).fetchall()
    return "\n".join(f"{r['role']}: {r['content']}" for r in rows)


def _system_content(conn, conversation_id):
    row = conn.execute(
        "SELECT content FROM messages WHERE conversation_id = ? AND role = 'system' LIMIT 1",
        (conversation_id,)
    ).fetchone()
    return row["content"] if row else ""


def _insert_message(conn, conversation_id, role, content):
    cur = conn.execute(
        "INSERT INTO messages (conversation_id, role, content) VALUES (?, ?, ?)",
        (conversation_id, role, content)
    )
    conn.commit()
    return cur.lastrowid


def _block_html(mode_key, message_id, text):
    meta = MODE_BY_KEY[mode_key]
    safe_text = str(escape(text or ""))
    return (
        f'<div class="ai-block" data-ai-block="1" data-mode="{mode_key}" data-msg-id="{message_id}" contenteditable="false">'
        f'<div class="ai-block-header">'
        f'<span class="ai-block-label">{meta["icon"]} {escape(meta["title"])}</span>'
        f'<button type="button" class="ai-block-remove" onclick="removeAiBlock(this)" title="Elimină din document">&times;</button>'
        f'</div>'
        f'<div class="math-content" data-math="{safe_text}">{safe_text}</div>'
        f'</div><p><br></p>'
    )


def _append_to_document(conn, document_id, block_html):
    doc = conn.execute("SELECT content FROM documents WHERE id = ?", (document_id,)).fetchone()
    new_content = (doc["content"] or "") + block_html
    conn.execute(
        "UPDATE documents SET content = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
        (new_content, document_id)
    )
    conn.commit()


def _serialize_message(row, block_html=None):
    return {
        "id": row["id"],
        "role": row["role"],
        "content": row["content"],
        "block_html": block_html,
    }


# =========================================================
# DOCUMENT CRUD
# =========================================================

@document_editor_bp.route("/editor")
def editor_home():
    if not _require_login():
        return redirect("/login")

    conn = get_db()
    latest = conn.execute(
        "SELECT id FROM documents WHERE user_id = ? ORDER BY updated_at DESC LIMIT 1",
        (session["user_id"],)
    ).fetchone()

    if latest:
        return redirect(f"/editor/{latest['id']}")

    cur = conn.execute(
        "INSERT INTO documents (user_id, title, content) VALUES (?, ?, ?)",
        (session["user_id"], "Document nou", "")
    )
    conn.commit()
    return redirect(f"/editor/{cur.lastrowid}")


@document_editor_bp.route("/editor/new", methods=["POST"])
def editor_new():
    if not _require_login():
        return redirect("/login")

    conn = get_db()
    cur = conn.execute(
        "INSERT INTO documents (user_id, title, content) VALUES (?, ?, ?)",
        (session["user_id"], "Document nou", "")
    )
    conn.commit()
    return redirect(f"/editor/{cur.lastrowid}")


@document_editor_bp.route("/editor/<int:document_id>")
def editor_view(document_id):
    if not _require_login():
        return redirect("/login")

    conn = get_db()
    document = _doc_or_none(conn, document_id, session["user_id"])
    if not document:
        return "Document inexistent", 404

    documents = conn.execute(
        "SELECT * FROM documents WHERE user_id = ? ORDER BY updated_at DESC",
        (session["user_id"],)
    ).fetchall()

    styles = conn.execute(
        "SELECT * FROM styles WHERE user_id = ? ORDER BY id DESC",
        (session["user_id"],)
    ).fetchall()

    mode_state = {}
    for m in MODES:
        conv = _mode_conversation(conn, document_id, m["key"], session["user_id"])
        if conv:
            messages = [_serialize_message(r) for r in _mode_messages(conn, conv["id"]) if r["role"] != "system"]
        else:
            messages = []
        mode_state[m["key"]] = {
            "conversation_id": conv["id"] if conv else None,
            "messages": messages,
        }

    return render_template(
        "editor.html",
        document=document,
        documents=documents,
        styles=styles,
        modes=MODES,
        mode_state=mode_state,
    )


@document_editor_bp.route("/editor/<int:document_id>/save", methods=["POST"])
def editor_save(document_id):
    if not _require_login():
        return jsonify({"ok": False, "error": "auth"}), 401

    conn = get_db()
    document = _doc_or_none(conn, document_id, session["user_id"])
    if not document:
        return jsonify({"ok": False, "error": "not_found"}), 404

    data = request.get_json(silent=True) or {}
    content = data.get("content", document["content"])

    conn.execute(
        "UPDATE documents SET content = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
        (content, document_id)
    )
    conn.commit()
    return jsonify({"ok": True, "updated_at": datetime.utcnow().isoformat()})


@document_editor_bp.route("/editor/<int:document_id>/rename", methods=["POST"])
def editor_rename(document_id):
    if not _require_login():
        return redirect("/login")

    conn = get_db()
    document = _doc_or_none(conn, document_id, session["user_id"])
    if not document:
        return "Document inexistent", 404

    new_title = (request.form.get("new_title") or "").strip()
    if new_title:
        conn.execute(
            "UPDATE documents SET title = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
            (new_title, document_id)
        )
        conn.commit()
    return redirect(f"/editor/{document_id}")


@document_editor_bp.route("/editor/<int:document_id>/delete", methods=["POST"])
def editor_delete(document_id):
    if not _require_login():
        return redirect("/login")

    conn = get_db()
    document = _doc_or_none(conn, document_id, session["user_id"])
    if not document:
        return "Document inexistent", 404

    convo_ids = [r["id"] for r in conn.execute(
        "SELECT id FROM conversations WHERE document_id = ?", (document_id,)
    ).fetchall()]
    for cid in convo_ids:
        conn.execute("DELETE FROM messages WHERE conversation_id = ?", (cid,))
    conn.execute("DELETE FROM conversations WHERE document_id = ?", (document_id,))
    conn.execute("DELETE FROM documents WHERE id = ? AND user_id = ?", (document_id, session["user_id"]))
    conn.commit()

    next_doc = conn.execute(
        "SELECT id FROM documents WHERE user_id = ? ORDER BY updated_at DESC LIMIT 1",
        (session["user_id"],)
    ).fetchone()

    if next_doc:
        return redirect(f"/editor/{next_doc['id']}")
    return redirect("/editor/new")


@document_editor_bp.route("/editor/<int:document_id>/open", methods=["POST"])
def editor_open_file(document_id):
    if not _require_login():
        return jsonify({"ok": False, "error": "auth"}), 401

    conn = get_db()
    document = _doc_or_none(conn, document_id, session["user_id"])
    if not document:
        return jsonify({"ok": False, "error": "not_found"}), 404

    uploaded = request.files.get("file")
    if not uploaded or uploaded.filename == "":
        return jsonify({"ok": False, "error": "Niciun fișier selectat"}), 400

    if not _allowed(uploaded.filename, MODEL_EXTENSIONS):
        return jsonify({"ok": False, "error": "Tip de fișier nepermis (.pdf, .docx, .txt)"}), 400

    path, filename = _save_upload(document_id, "open", uploaded)
    html_fragment = _open_file_to_html(path)

    return jsonify({"ok": True, "html": html_fragment, "filename": filename})


_DATA_URL_RE = re.compile(r"^data:image/(png|jpe?g);base64,(.+)$", re.DOTALL)


def _latex_blocks_from_transcription(text):
    """Split a transcribe_handwriting() response (one or more $$...$$ blocks)
    into individual LaTeX strings, stripped of the delimiters."""
    blocks = re.findall(r"\$\$(.*?)\$\$", text or "", re.DOTALL)
    blocks = [b.strip() for b in blocks if b.strip()]
    if blocks:
        return blocks
    stripped = (text or "").strip()
    if not stripped or stripped.lower().startswith("eroare"):
        return []
    return [stripped]


@document_editor_bp.route("/editor/<int:document_id>/ink-transcribe", methods=["POST"])
def editor_ink_transcribe(document_id):
    """Ink Equation: OCRs a hand-drawn canvas sketch into LaTeX using the
    same handwriting-transcription pipeline as mode4, so it renders back
    into the document as a real (KaTeX) equation instead of a static image."""
    if not _require_login():
        return jsonify({"ok": False, "error": "auth"}), 401

    conn = get_db()
    document = _doc_or_none(conn, document_id, session["user_id"])
    if not document:
        return jsonify({"ok": False, "error": "not_found"}), 404

    data = request.get_json(silent=True) or {}
    match = _DATA_URL_RE.match(data.get("image", ""))
    if not match:
        return jsonify({"ok": False, "error": "invalid_image"}), 400

    extension = "png" if match.group(1) == "png" else "jpg"
    try:
        raw = base64.b64decode(match.group(2))
    except (ValueError, TypeError):
        return jsonify({"ok": False, "error": "invalid_image"}), 400

    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    path = os.path.join(UPLOAD_FOLDER, f"ink_{document_id}_{uuid.uuid4().hex}.{extension}")
    with open(path, "wb") as f:
        f.write(raw)

    try:
        transcription = ai.transcribe_handwriting(path)
    finally:
        try:
            os.remove(path)
        except OSError:
            pass

    blocks = _latex_blocks_from_transcription(transcription)
    if not blocks:
        return jsonify({"ok": False, "error": "empty_transcription"}), 400

    return jsonify({"ok": True, "latex_blocks": blocks})


def _open_file_to_html(path):
    extension = os.path.splitext(path)[1].lower()

    if extension == ".txt":
        with open(path, "r", encoding="utf-8") as f:
            lines = f.read().splitlines()
        return "".join(f"<p>{escape(line) if line.strip() else '<br>'}</p>" for line in lines) or "<p><br></p>"

    if extension == ".docx":
        doc = DocxDocument(path)
        parts = []
        for p in doc.paragraphs:
            if not p.text.strip():
                parts.append("<p><br></p>")
                continue
            style_name = (p.style.name or "").lower() if p.style else ""
            tag = "p"
            if "heading 1" in style_name or "title" in style_name:
                tag = "h1"
            elif "heading 2" in style_name:
                tag = "h2"
            elif "heading 3" in style_name:
                tag = "h3"
            runs_html = []
            for run in p.runs:
                text = str(escape(run.text))
                if not text:
                    continue
                if run.bold:
                    text = f"<strong>{text}</strong>"
                if run.italic:
                    text = f"<em>{text}</em>"
                if run.underline:
                    text = f"<u>{text}</u>"
                runs_html.append(text)
            parts.append(f"<{tag}>{''.join(runs_html) or escape(p.text)}</{tag}>")
        return "".join(parts) or "<p><br></p>"

    if extension == ".pdf":
        text = _file_read(path)
        lines = text.splitlines()
        return "".join(f"<p>{escape(line) if line.strip() else '<br>'}</p>" for line in lines) or "<p><br></p>"

    return "<p><br></p>"


@document_editor_bp.route("/editor/<int:document_id>/print")
def editor_print(document_id):
    if not _require_login():
        return redirect("/login")

    conn = get_db()
    document = _doc_or_none(conn, document_id, session["user_id"])
    if not document:
        return "Document inexistent", 404

    return render_template("editor_print.html", document=document)


@document_editor_bp.route("/editor/<int:document_id>/download")
def editor_download(document_id):
    if not _require_login():
        return redirect("/login")

    conn = get_db()
    document = _doc_or_none(conn, document_id, session["user_id"])
    if not document:
        return "Document inexistent", 404

    docx_doc = _html_to_docx(document["content"] or "")
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    filename = f"document_{document_id}.docx"
    path = os.path.join(UPLOAD_FOLDER, filename)
    docx_doc.save(path)

    safe_title = "".join(c for c in (document["title"] or "Document") if c.isalnum() or c in " -_").strip() or "Document"
    return send_file(
        path,
        as_attachment=True,
        download_name=f"{safe_title}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


class _BlockHTMLToDocx(HTMLParser):
    """Minimal HTML -> docx converter for the editor's contenteditable output.
    Handles paragraphs/headings with bold/italic/underline runs, lists, and
    math-content blocks (inserted as plain LaTeX text — no OMML conversion)."""

    BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "li"}

    def __init__(self, docx_doc):
        super().__init__()
        self.doc = docx_doc
        self.paragraph = None
        self.bold = 0
        self.italic = 0
        self.underline = 0
        self.heading_level = None
        self.skip_depth = 0

    def _ensure_paragraph(self):
        if self.paragraph is None:
            if self.heading_level:
                self.paragraph = self.doc.add_heading(level=self.heading_level)
            else:
                self.paragraph = self.doc.add_paragraph()
        return self.paragraph

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)
        if tag == "button":
            self.skip_depth += 1
            return
        if tag in ("strong", "b"):
            self.bold += 1
        elif tag in ("em", "i"):
            self.italic += 1
        elif tag == "u":
            self.underline += 1
        elif tag == "br":
            self._ensure_paragraph().add_run().add_break()
        elif tag == "h1":
            self.heading_level = 1
            self.paragraph = None
        elif tag == "h2":
            self.heading_level = 2
            self.paragraph = None
        elif tag == "h3":
            self.heading_level = 3
            self.paragraph = None
        elif tag in self.BLOCK_TAGS:
            self.paragraph = None

    def handle_endtag(self, tag):
        if tag == "button":
            self.skip_depth = max(0, self.skip_depth - 1)
            return
        if tag in ("strong", "b"):
            self.bold = max(0, self.bold - 1)
        elif tag in ("em", "i"):
            self.italic = max(0, self.italic - 1)
        elif tag == "u":
            self.underline = max(0, self.underline - 1)
        elif tag in ("h1", "h2", "h3"):
            self.heading_level = None
            self.paragraph = None
        elif tag in self.BLOCK_TAGS:
            self.paragraph = None

    def handle_data(self, data):
        if self.skip_depth:
            return
        text = data.strip("\n")
        if not text.strip():
            return
        run = self._ensure_paragraph().add_run(text)
        run.bold = bool(self.bold)
        run.italic = bool(self.italic)
        run.underline = bool(self.underline)


def _html_to_docx(html_content):
    docx_doc = DocxDocument()
    parser = _BlockHTMLToDocx(docx_doc)
    parser.feed(html_content or "")
    if not docx_doc.paragraphs:
        docx_doc.add_paragraph("")
    return docx_doc


# =========================================================
# MODE PANELS (AI writes into the document)
# =========================================================

def _run_generation(mode_key, action_type, conn, conversation, form, files, is_start):
    """Returns (raw_text, updated_title_or_None)."""

    if mode_key == "mode1":
        prompt = form.get("prompt", "").strip()
        history = _history_text(conn, conversation["id"])
        full_prompt = f"{history}\nuser: {prompt}" if history else prompt
        return ai.assistant(full_prompt), None

    if mode_key == "mode2":
        style_desc = conversation["style_description"] if "style_description" in conversation.keys() else None
        style_desc = style_desc or "Stil implicit"
        school_class = conversation["school_class"] or "—"
        bac = conversation["bac"] or "—"
        prompt = form.get("prompt", "").strip()

        if action_type == "generate":
            try:
                n = max(1, min(int(form.get("exercise_count", "5")), 50))
            except (TypeError, ValueError):
                n = 5
            return adjacent.generate_exercises_free(n, style_desc, school_class, bac), None

        if action_type == "solve":
            target = prompt or _last_user_content(conn, conversation["id"])
            return adjacent.solve_step_by_step(target, style_desc, school_class, bac), None

        return ai.translate_math(prompt, style_desc, school_class, bac), None

    if mode_key in ("mode3", "mode4"):
        if is_start:
            uploaded = files.get("file")
            if not uploaded or uploaded.filename == "":
                raise ValueError("Niciun fișier încărcat")
            allowed = MODEL_EXTENSIONS if mode_key == "mode3" else SCAN_EXTENSIONS
            if not _allowed(uploaded.filename, allowed):
                raise ValueError("Tip de fișier nepermis")

            path, filename = _save_upload(conversation["document_id"], mode_key, uploaded)

            if mode_key == "mode3":
                raw_content = _file_read(path)
                latex_content = ai.convert_file_to_latex(raw_content, filepath=path)
            else:
                latex_content = ai.transcribe_handwriting(path)

            _insert_message(conn, conversation["id"], "system", latex_content)

            prompt = form.get("prompt", "").strip()
            differences = form.get("differences", "").strip()
            user_instruction = f"Instrucțiuni: {prompt}\nDiferențe față de model: {differences}".strip()

            generated = ai.generate_from_model(
                model_content=latex_content,
                conversation_history="",
                user_prompt=user_instruction
            )
            return generated, filename

        # follow-up chat message
        prompt = form.get("prompt", "").strip()
        model_content = _system_content(conn, conversation["id"])
        history = _history_text(conn, conversation["id"])
        return ai.generate_from_model(
            model_content=model_content,
            conversation_history=history,
            user_prompt=prompt
        ), None

    if mode_key == "mode5":
        if is_start:
            lessons = form.get("lessons", "").strip()
            avoid = form.get("avoid", "").strip()
            return ai.bac_generator(lessons, avoid), None

        prompt = form.get("prompt", "").strip()

        quick_prompts = {
            "solve_bac": "Rezolvă complet, pas cu pas, următoarea variantă BAC. Include baremul, explicațiile și răspunsurile finale:\n\n",
            "generate_similar_bac": "Generează o variantă BAC similară cu următoarea. Păstrează structura oficială, nivelul de dificultate, dar schimbă exercițiile și valorile numerice:\n\n",
            "generate_row_2_bac": "Generează rândul 2 pentru următoarea variantă BAC. Păstrează aceeași structură și nivel de dificultate, dar creează exerciții diferite:\n\n",
        }
        if action_type in quick_prompts:
            last_bac = _last_assistant_content(conn, conversation["id"])
            prompt = quick_prompts[action_type] + last_bac

        model_content = _last_assistant_content(conn, conversation["id"])
        history = _history_text(conn, conversation["id"])
        return ai.generate_from_model(
            model_content=model_content,
            conversation_history=history,
            user_prompt=prompt
        ), None

    raise ValueError("Mod necunoscut")


@document_editor_bp.route("/editor/<int:document_id>/mode/<mode_key>/start", methods=["POST"])
def mode_start(document_id, mode_key):
    if not _require_login():
        return jsonify({"ok": False, "error": "auth"}), 401
    if mode_key not in MODE_BY_KEY:
        return jsonify({"ok": False, "error": "unknown_mode"}), 404

    conn = get_db()
    document = _doc_or_none(conn, document_id, session["user_id"])
    if not document:
        return jsonify({"ok": False, "error": "not_found"}), 404

    existing = _mode_conversation(conn, document_id, mode_key, session["user_id"])
    if existing:
        return jsonify({"ok": False, "error": "already_started", "conversation_id": existing["id"]}), 400

    style_id = request.form.get("style") or None
    school_class = request.form.get("school_class") or None
    bac = request.form.get("bac") or None

    cur = conn.execute(
        """
        INSERT INTO conversations (user_id, mode, title, document_id, style_id, school_class, bac)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (session["user_id"], mode_key, MODE_BY_KEY[mode_key]["title"], document_id, style_id, school_class, bac)
    )
    conn.commit()
    conversation_id = cur.lastrowid

    conversation = conn.execute(
        """
        SELECT conversations.*, styles.style_description
        FROM conversations LEFT JOIN styles ON conversations.style_id = styles.id
        WHERE conversations.id = ?
        """,
        (conversation_id,)
    ).fetchone()

    try:
        if mode_key == "mode1":
            intro = (
                "Sunt un asistent virtual specializat în matematică. Pot ajuta cu explicații, "
                "exerciții, predare și gestionarea elevilor. Cu ce te pot ajuta astăzi?"
            )
            msg_id = _insert_message(conn, conversation_id, "assistant", intro)
            block = _block_html(mode_key, msg_id, intro)
            return jsonify({
                "ok": True, "conversation_id": conversation_id,
                "message": _serialize_message(conn.execute("SELECT * FROM messages WHERE id=?", (msg_id,)).fetchone(), block),
                "auto_insert": False,
            })

        if mode_key == "mode2":
            style_desc = conversation["style_description"] or "Stil implicit"
            intro = (
                f"Stilul tău activ este:\n- {style_desc}\n- Clasa: {school_class or '—'}\n- BAC: {bac or '—'}\n\n"
                f"Scrie exercițiile în limbaj natural, iar eu le voi transforma în limbaj matematic."
            )
            msg_id = _insert_message(conn, conversation_id, "assistant", intro)
            return jsonify({
                "ok": True, "conversation_id": conversation_id,
                "message": _serialize_message(conn.execute("SELECT * FROM messages WHERE id=?", (msg_id,)).fetchone()),
                "auto_insert": False,
            })

        generated, extra_title = _run_generation(
            mode_key, "normal", conn, dict(conversation), request.form, request.files, True
        )
    except ValueError as e:
        conn.execute("DELETE FROM conversations WHERE id = ?", (conversation_id,))
        conn.commit()
        return jsonify({"ok": False, "error": str(e)}), 400

    _insert_message(conn, conversation_id, "user", "Generare inițială")
    msg_id = _insert_message(conn, conversation_id, "assistant", generated)

    if extra_title:
        conn.execute("UPDATE conversations SET title = ? WHERE id = ?", (extra_title, conversation_id))
        conn.commit()

    block = _block_html(mode_key, msg_id, generated)
    if MODE_BY_KEY[mode_key]["auto_insert"]:
        _append_to_document(conn, document_id, block)

    return jsonify({
        "ok": True,
        "conversation_id": conversation_id,
        "message": _serialize_message(conn.execute("SELECT * FROM messages WHERE id=?", (msg_id,)).fetchone(), block),
        "auto_insert": MODE_BY_KEY[mode_key]["auto_insert"],
    })


@document_editor_bp.route("/editor/<int:document_id>/mode/<mode_key>/message", methods=["POST"])
def mode_message(document_id, mode_key):
    if not _require_login():
        return jsonify({"ok": False, "error": "auth"}), 401
    if mode_key not in MODE_BY_KEY:
        return jsonify({"ok": False, "error": "unknown_mode"}), 404

    conn = get_db()
    document = _doc_or_none(conn, document_id, session["user_id"])
    if not document:
        return jsonify({"ok": False, "error": "not_found"}), 404

    conversation = conn.execute(
        """
        SELECT conversations.*, styles.style_description
        FROM conversations LEFT JOIN styles ON conversations.style_id = styles.id
        WHERE conversations.document_id = ? AND conversations.mode = ? AND conversations.user_id = ?
        """,
        (document_id, mode_key, session["user_id"])
    ).fetchone()

    if not conversation:
        return jsonify({"ok": False, "error": "not_started"}), 400

    action_type = request.form.get("action_type", "normal")
    prompt = request.form.get("prompt", "").strip()

    if not prompt and action_type == "normal":
        return jsonify({"ok": False, "error": "empty_prompt"}), 400

    if prompt:
        _insert_message(conn, conversation["id"], "user", prompt)

    try:
        generated, _ = _run_generation(mode_key, action_type, conn, dict(conversation), request.form, request.files, False)
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 400

    msg_id = _insert_message(conn, conversation["id"], "assistant", generated)
    block = _block_html(mode_key, msg_id, generated)

    if MODE_BY_KEY[mode_key]["auto_insert"]:
        _append_to_document(conn, document_id, block)

    return jsonify({
        "ok": True,
        "conversation_id": conversation["id"],
        "message": _serialize_message(conn.execute("SELECT * FROM messages WHERE id=?", (msg_id,)).fetchone(), block),
        "auto_insert": MODE_BY_KEY[mode_key]["auto_insert"],
    })


@document_editor_bp.route("/editor/<int:document_id>/mode/<mode_key>/insert", methods=["POST"])
def mode_insert(document_id, mode_key):
    """Manual insert-to-document for panels that don't auto-insert (mode1)."""
    if not _require_login():
        return jsonify({"ok": False, "error": "auth"}), 401

    conn = get_db()
    document = _doc_or_none(conn, document_id, session["user_id"])
    if not document:
        return jsonify({"ok": False, "error": "not_found"}), 404

    data = request.get_json(silent=True) or {}
    message_id = data.get("message_id")
    message = conn.execute("SELECT * FROM messages WHERE id = ?", (message_id,)).fetchone()
    if not message:
        return jsonify({"ok": False, "error": "not_found"}), 404

    block = _block_html(mode_key, message["id"], message["content"])
    _append_to_document(conn, document_id, block)
    return jsonify({"ok": True, "block_html": block})
